"""Helper script: create a chapter .docx from text content."""
import docx
from docx.shared import Pt, Inches
import sys

def make_doc(filepath, text):
    doc = docx.Document()

    # Set reasonable margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Remove the default empty paragraph Word always inserts
    for elem in list(doc.element.body):
        doc.element.body.remove(elem)

    blocks = [b for b in text.strip().split('\n\n') if b.strip()]

    for block in blocks:
        stripped = block.strip()
        lines = stripped.split('\n')

        if stripped == '***':
            p = doc.add_paragraph('***')
        elif stripped.startswith('---'):
            # System block: each line gets its own paragraph
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
        else:
            p = doc.add_paragraph(stripped)

        if stripped != '***' and not stripped.startswith('---'):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

    doc.save(filepath)
    print(f'Saved: {filepath}')

if __name__ == '__main__':
    # Can be called with: python write_chapter.py <filepath> <textfile>
    if len(sys.argv) == 3:
        with open(sys.argv[2], encoding='utf-8') as f:
            text = f.read()
        make_doc(sys.argv[1], text)
